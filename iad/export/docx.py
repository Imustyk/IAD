"""DOCX report generation using python-docx."""
from __future__ import annotations

from pathlib import Path

from iad.core.exceptions import ExportError
from iad.core.logging import get_logger
from iad.export.models import AnalyticsReport, ExportFormat, ExportResult

logger = get_logger("iad.export.docx")


def build_docx_report(report: AnalyticsReport, destination: Path) -> ExportResult:
    """Render an analytics report as a Word document."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt
    except ImportError as exc:
        raise ExportError(
            "DOCX export requires python-docx. Install with: pip install python-docx",
            code="docx_missing",
        ) from exc

    destination = Path(destination)
    destination.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    title = document.add_heading(report.title, level=0)
    title.runs[0].font.size = Pt(22)

    if report.subtitle:
        document.add_paragraph(report.subtitle)
    if report.generated_at_iso:
        document.add_paragraph(f"Generated: {report.generated_at_iso}")

    for line in report.summary_lines():
        document.add_paragraph(line)

    if report.metrics:
        document.add_heading("Model metrics", level=2)
        table = document.add_table(rows=1, cols=2)
        hdr = table.rows[0].cells
        hdr[0].text = "Metric"
        hdr[1].text = "Value"
        for key, value in report.metrics.items():
            row = table.add_row().cells
            row[0].text = str(key)
            row[1].text = f"{value:.6g}" if isinstance(value, float) else str(value)

    for section in report.sections:
        document.add_heading(section.title, level=2)
        document.add_paragraph(section.body)
        for bullet in section.bullets:
            document.add_paragraph(bullet, style="List Bullet")
        if section.metrics:
            stable = document.add_table(rows=1, cols=2)
            stable.rows[0].cells[0].text = "Metric"
            stable.rows[0].cells[1].text = "Value"
            for key, value in section.metrics.items():
                cells = stable.add_row().cells
                cells[0].text = str(key)
                if isinstance(value, float):
                    cells[1].text = f"{value:.6g}"
                else:
                    cells[1].text = str(value)

    for chart_path in report.chart_paths:
        if chart_path.exists() and chart_path.suffix.lower() in {".png", ".jpg", ".jpeg"}:
            document.add_picture(str(chart_path), width=Inches(6))

    try:
        document.save(str(destination))
    except Exception as exc:
        logger.exception("DOCX build failed")
        raise ExportError(f"DOCX generation failed: {exc}", code="docx_build_failed") from exc

    logger.info("DOCX exported", extra={"ctx_path": str(destination)})
    return ExportResult(
        format=ExportFormat.DOCX,
        path=destination,
        size_bytes=destination.stat().st_size,
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
